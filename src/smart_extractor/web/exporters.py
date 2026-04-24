"""
任务结果导出工具。
"""

from __future__ import annotations

import json
from io import BytesIO
from typing import Any

from docx import Document
from openpyxl import Workbook


def _stringify(value: Any) -> str:
    if value in (None, "", [], {}):
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


def build_task_markdown(detail: dict[str, Any]) -> str:
    lines = [
        f"# Smart Extractor 任务报告 - {detail.get('task_id', '-')}",
        "",
        "## 任务概览",
        "",
        f"- URL: {detail.get('url', '-')}",
        f"- 状态: {detail.get('status', '-')}",
        f"- 存储格式: {detail.get('storage_format', '-')}",
        f"- 域名: {detail.get('domain', '-')}",
        f"- 创建时间: {detail.get('created_at', '-')}",
        f"- 完成时间: {detail.get('completed_at', '-')}",
        f"- 耗时: {round(float(detail.get('elapsed_ms', 0) or 0))} ms",
        f"- 质量分: {float(detail.get('quality_score', 0) or 0):.1%}",
        "",
    ]
    comparison = detail.get("comparison") or {}
    if comparison.get("has_previous"):
        lines.extend(
            [
                "## 变化告警",
                "",
                f"- 是否有变化: {'是' if comparison.get('changed') else '否'}",
                f"- 变化字段数: {comparison.get('changed_fields_count', 0)}",
            ]
        )
        if comparison.get("impact_summary"):
            lines.append(f"- 影响摘要: {comparison.get('impact_summary')}")
        if comparison.get("summary_lines"):
            lines.extend(["", "### 变化摘要"])
            lines.extend([f"- {item}" for item in comparison["summary_lines"]])
        if comparison.get("suggested_actions"):
            lines.extend(["", "### 建议动作"])
            lines.extend([f"- {item}" for item in comparison["suggested_actions"]])
        lines.append("")

    data = detail.get("data") or {}
    if data:
        lines.extend(["## 提取结果", ""])
        for key, value in data.items():
            lines.append(f"- {key}: {_stringify(value)}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_task_docx(detail: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading(f"Smart Extractor 任务报告 - {detail.get('task_id', '-')}", 0)

    overview = document.add_table(rows=0, cols=2)
    overview.style = "Table Grid"
    rows = [
        ("任务编号", detail.get("task_id", "-")),
        ("URL", detail.get("url", "-")),
        ("任务状态", detail.get("status", "-")),
        ("存储格式", detail.get("storage_format", "-")),
        ("页面域名", detail.get("domain", "-")),
        ("创建时间", detail.get("created_at", "-")),
        ("完成时间", detail.get("completed_at", "-")),
        ("耗时", f"{round(float(detail.get('elapsed_ms', 0) or 0))} ms"),
        ("质量分", f"{float(detail.get('quality_score', 0) or 0):.1%}"),
    ]
    data = detail.get("data") or {}
    usage_payload = data.get("_llm_usage") if isinstance(data, dict) else {}
    if isinstance(usage_payload, dict) and usage_payload:
        rows.extend(
            [
                ("LLM 调用次数", int(usage_payload.get("total_calls", 0) or 0)),
                ("Prompt Tokens", int(usage_payload.get("prompt_tokens", 0) or 0)),
                ("Completion Tokens", int(usage_payload.get("completion_tokens", 0) or 0)),
                ("估算成本", f"${float(usage_payload.get('estimated_cost_usd', 0.0) or 0.0):.6f}"),
            ]
        )
    for label, value in rows:
        cells = overview.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    comparison = detail.get("comparison") or {}
    if comparison.get("has_previous"):
        document.add_heading("变化告警", level=1)
        document.add_paragraph(
            "检测结果："
            + (
                f"发现 {comparison.get('changed_fields_count', 0)} 个字段变化"
                if comparison.get("changed")
                else "与上一次成功结果相比没有字段变化"
            )
        )
        impact_summary = str(comparison.get("impact_summary") or "").strip()
        if impact_summary:
            document.add_paragraph("影响摘要：" + impact_summary)
        if comparison.get("summary_lines"):
            for line in comparison["summary_lines"]:
                document.add_paragraph(str(line), style="List Bullet")
        if comparison.get("suggested_actions"):
            document.add_paragraph("建议动作：")
            for line in comparison["suggested_actions"]:
                document.add_paragraph(str(line), style="List Bullet")

    if data:
        document.add_heading("提取结果", level=1)
        for key, value in data.items():
            document.add_paragraph(f"{key}：{_stringify(value)}")

    formatted_text = str(data.get("formatted_text") or "").strip() if isinstance(data, dict) else ""
    if formatted_text:
        document.add_heading("润色结果", level=1)
        document.add_paragraph(formatted_text)

    history = detail.get("recent_history") or []
    if history:
        document.add_heading("最近历史", level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "任务编号"
        header[1].text = "状态"
        header[2].text = "质量分"
        header[3].text = "时间"
        for item in history:
            cells = table.add_row().cells
            cells[0].text = str(item.get("task_id", "-"))
            cells[1].text = str(item.get("status", "-"))
            cells[2].text = f"{float(item.get('quality_score', 0) or 0):.1%}"
            cells[3].text = str(item.get("created_at", "-"))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_task_xlsx(detail: dict[str, Any]) -> bytes:
    workbook = Workbook()
    overview = workbook.active
    overview.title = "任务概览"
    overview.append(["字段", "值"])
    overview_rows = [
        ("任务编号", detail.get("task_id", "-")),
        ("URL", detail.get("url", "-")),
        ("任务状态", detail.get("status", "-")),
        ("存储格式", detail.get("storage_format", "-")),
        ("页面域名", detail.get("domain", "-")),
        ("创建时间", detail.get("created_at", "-")),
        ("完成时间", detail.get("completed_at", "-")),
        ("耗时(ms)", round(float(detail.get("elapsed_ms", 0) or 0))),
        ("质量分", float(detail.get("quality_score", 0) or 0)),
    ]
    usage_payload = (
        detail.get("data", {}).get("_llm_usage")
        if isinstance(detail.get("data"), dict)
        else {}
    )
    if isinstance(usage_payload, dict) and usage_payload:
        overview_rows.extend(
            [
                ("LLM 调用次数", int(usage_payload.get("total_calls", 0) or 0)),
                ("Prompt Tokens", int(usage_payload.get("prompt_tokens", 0) or 0)),
                ("Completion Tokens", int(usage_payload.get("completion_tokens", 0) or 0)),
                ("估算成本(USD)", float(usage_payload.get("estimated_cost_usd", 0.0) or 0.0)),
            ]
        )
    for row in overview_rows:
        overview.append(row)

    result_sheet = workbook.create_sheet("提取结果")
    result_sheet.append(["字段", "值"])
    data = detail.get("data") or {}
    for key, value in data.items():
        result_sheet.append([key, _stringify(value)])

    comparison_sheet = workbook.create_sheet("变化明细")
    comparison_sheet.append(["字段", "展示名", "变化类型", "之前", "现在", "摘要"])
    comparison = detail.get("comparison") or {}
    for item in comparison.get("changed_fields", []) or []:
        comparison_sheet.append(
            [
                item.get("field", ""),
                item.get("label", ""),
                item.get("change_type", ""),
                _stringify(item.get("before")),
                _stringify(item.get("after")),
                item.get("summary", ""),
            ]
        )

    action_sheet = workbook.create_sheet("变化建议")
    action_sheet.append(["类型", "内容"])
    if comparison.get("impact_summary"):
        action_sheet.append(["影响摘要", comparison.get("impact_summary", "")])
    for item in comparison.get("suggested_actions", []) or []:
        action_sheet.append(["建议动作", item])

    history_sheet = workbook.create_sheet("历史记录")
    history_sheet.append(["任务编号", "状态", "质量分", "创建时间", "完成时间"])
    for item in detail.get("recent_history", []) or []:
        history_sheet.append(
            [
                item.get("task_id", ""),
                item.get("status", ""),
                float(item.get("quality_score", 0) or 0),
                item.get("created_at", ""),
                item.get("completed_at", ""),
            ]
        )

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()
